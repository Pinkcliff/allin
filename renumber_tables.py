# -*- coding: utf-8 -*-
"""重新排列表格序号 - 每个表格独立从1开始"""

from docx import Document

doc = Document(r'F:\A-User\cliff\allin\软著申请材料\软著申请材料_完整版_已修改_新版.docx')

# 需要重新编号的表格索引（有序号列的表格）
numbered_tables = [10, 11, 12, 13, 14, 15, 16, 17, 18]

# 重新编号每个表格，每个表格独立从1开始
for table_idx in numbered_tables:
    if table_idx >= len(doc.tables):
        continue

    table = doc.tables[table_idx]

    # 检查第一列是否是"序号"
    if len(table.rows) > 0 and len(table.rows[0].cells) > 0:
        first_cell = table.rows[0].cells[0]
        if '序号' in first_cell.text:
            # 从第2行开始重新编号（第1行是表头），每个表格独立从1开始
            local_number = 1
            for row_idx in range(1, len(table.rows)):
                row = table.rows[row_idx]
                # 更新第一列的序号
                row.cells[0].text = str(local_number)
                local_number += 1

            print(f'表格 {table_idx}: 已重新编号 1-{len(table.rows)-1}')

# 保存文档
doc.save(r'F:\A-User\cliff\allin\软著申请材料\软著申请材料_完整版_已修改_新版.docx')
print('\\n序号重新编排完成！每个表格独立从1开始编号。')
