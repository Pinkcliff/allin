# -*- coding: utf-8 -*-
"""处理软著申请材料文档，删除CFD和动态控制相关内容"""

from docx import Document
import sys

def process_document():
    # 读取原始文档
    doc = Document(r'F:\A-User\cliff\allin\软著申请材料\软著申请材料_完整版_已修改.docx')

    # 需要删除的段落内容（完整匹配）
    delete_texts = [
        'CFD前处理 - 场景生成、网格生成、边界条件',
        '• CFD前处理模块',
        '3.2.5 CFD前处理模块',
        '• 表达式解析：支持自定义数学表达式定义风速分布',
        '• 表达式解析 - 动态数学表达式解析',
        '• 时间轴控制：支持风场随时间变化的动画编辑',
        '• 自定义表达式',
        '表达式解析引擎：支持复杂数学表达式的动态解析和计算',
        '3. 自定义风场表达式：支持用户通过数学表达式定义任意风场分布',
        '  - 支持CSV格式控制序列导入和回放',
        '• 导出控制序列：导出到硬件控制系统',
        '33 | CSV控制序列 | 加载和回放CSV控制文件',
        '场景生成 | 根据风场配置生成CFD场景',
        '网格生成 | 自动生成计算域网格',
        '网格密度配置 | 可配置网格分辨率',
        '边界条件设置 | 自动设置入口/出口边界',
        '风扇模型映射 | 风扇阵列映射到CFD边界',
        'VTM文件导出 | 导出VTK多块数据集',
        '场景预览 | 预览生成的CFD场景',
        '16 | 自定义表达式 | 支持复杂数学表达式定义风速',
        '17 | 表达式解析器 | 动态解析和计算数学表达式',
        '19 | 时间轴编辑 | 支持风场随时间变化的动画',
        '20 | 关键帧管理 | 添加、编辑、删除关键帧',
        '21 | 插值计算 | 关键帧间自动插值',
        '24 | 导出控制序列 | 导出到硬件控制系统',
    ]

    # 需要替换的内容（旧 -> 新）
    replacements = [
        ('风扇阵列控制、环境参数监测、动作捕捉数据采集、CFD前处理等多个功能模块',
         '风扇阵列控制、环境参数监测、动作捕捉数据采集等多个功能模块'),
        ('集成了风扇阵列控制、环境参数监测、动作捕捉数据采集、CFD前处理等多个功能模块',
         '集成了风扇阵列控制、环境参数监测、动作捕捉数据采集等多个功能模块'),
        ('5. CFD前处理：生成风场测试场景配置文件', ''),
        ('一体化融合平台：首次将风场控制、数据采集、动捕、CFD等功能集成于统一平台',
         '一体化融合平台：首次将风场控制、数据采集、动捕等功能集成于统一平台'),
        ('• 风场编辑器 - 3D可视化编辑、风速区域、自定义表达式、模板库',
         '• 风场编辑器 - 3D可视化编辑、风速区域、模板库'),
        ('• 模板库管理：预设多种常用风场模式（层流、湍流、梯度风等）',
         '• 模板库管理：预设多种常用风场模式（层流、湍流、梯度风等）\n• 数据库配置导入：支持从数据库导入风扇转速配置列表'),
        ('  - 支持CSV格式控制序列导入和回放', '  - 支持从数据库导入风扇转速配置列表'),
        ('1. 风扇阵列控制：配置1600台风扇的静态转速参数',
         '1. 风扇阵列控制：从数据库导入1600台风扇的静态转速配置列表'),
        ('3.2.2 风场编辑器\n• 3D可视化编辑：在3D场景中布置风速区域和目标点\n• 表达式解析：支持自定义数学表达式定义风速分布\n• 模板库管理：预设多种常用风场模式（层流、湍流、梯度风等）\n• 时间轴控制：支持风场随时间变化的动画编辑\n• 实时预览：实时查看风场矢量和风速分布',
         '3.2.2 风场编辑器\n• 3D可视化编辑：在3D场景中布置风速区域和目标点\n• 模板库管理：预设多种常用风场模式（层流、湍流、梯度风等）\n• 数据库配置导入：支持从数据库导入风扇转速配置列表\n• 实时预览：实时查看风场矢量和风速分布'),
        ('3.2.5 CFD前处理模块\n• 场景生成：根据风场配置生成CFD场景文件\n• 网格生成：自动生成计算域网格\n• 边界条件设置：自动设置入口、出口、壁面等边界条件\n• 风扇模型映射：将风扇阵列映射到CFD边界条件',
         ''),
        ('自定义表达式引擎，灵活定义风场参数', '数据库配置管理，灵活导入风扇转速列表'),
        ('2.6 CFD前处理模块', ''),
        ('5. CFD前处理：生成风场测试场景配置文件', ''),
    ]

    # 收集需要删除的段落索引
    paragraphs_to_delete = []

    # 先处理删除
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        for delete_text in delete_texts:
            if delete_text in text:
                paragraphs_to_delete.append(i)
                break

    # 从后往前删除段落
    for idx in sorted(set(paragraphs_to_delete), reverse=True):
        if idx < len(doc.paragraphs):
            p = doc.paragraphs[idx]._element
            p.getparent().remove(p)

    # 处理替换
    for i, para in enumerate(doc.paragraphs):
        for old_text, new_text in replacements:
            if old_text in para.text:
                # 清空段落并重新添加文本
                para.text = para.text.replace(old_text, new_text)

    # 处理表格
    for table in doc.tables:
        # 收集需要删除的行
        rows_to_delete = []
        for row_idx, row in enumerate(table.rows):
            row_text = ' '.join([cell.text.strip() for cell in row.cells])
            for delete_text in delete_texts:
                if delete_text in row_text:
                    rows_to_delete.append(row_idx)
                    break

        # 删除表格行
        for row_idx in sorted(set(rows_to_delete), reverse=True):
            if row_idx < len(table.rows):
                tr = table.rows[row_idx]._element
                tr.getparent().remove(tr)

        # 处理表格中的替换
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for old_text, new_text in replacements:
                        if old_text in para.text:
                            para.text = para.text.replace(old_text, new_text)

    # 保存修改后的文档
    output_path = r'F:\A-User\cliff\allin\软著申请材料\软著申请材料_完整版_已修改_新版.docx'
    doc.save(output_path)
    print(f'文档处理完成！')
    print(f'删除了 {len(set(paragraphs_to_delete))} 个段落')
    print(f'保存到: {output_path}')

if __name__ == '__main__':
    process_document()
