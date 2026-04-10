# -*- coding: utf-8 -*-
"""
合并软著申请材料的Word文档
将多个Word文档合并为一个完整文档
"""
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def copy_table(source_table, target_doc):
    """复制表格到目标文档"""
    table = target_doc.add_table(rows=len(source_table.rows), cols=len(source_table.columns))
    table.style = source_table.style

    for i, row in enumerate(source_table.rows):
        for j, cell in enumerate(row.cells):
            table.rows[i].cells[j].text = cell.text
            # 复制格式
            for paragraph in table.rows[i].cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    run.font.size = Pt(11)

    return table


def copy_paragraph(source_paragraph, target_doc, keep_style=True):
    """复制段落到目标文档"""
    if source_paragraph.text.strip() == '':
        return target_doc.add_paragraph()

    p = target_doc.add_paragraph()
    p.alignment = source_paragraph.alignment

    for run in source_paragraph.runs:
        new_run = p.add_run(run.text)
        # 复制字体格式
        new_run.font.name = '宋体'
        new_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        new_run.font.size = run.font.size or Pt(12)
        new_run.font.bold = run.font.bold
        new_run.font.italic = run.font.italic

    return p


def set_chinese_font(run, font_name='宋体', font_size=12, bold=False):
    """设置中文字体"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    return run


def add_page_break_with_title(doc, title):
    """添加分页符和标题"""
    doc.add_page_break()
    h = doc.add_heading(title, level=1)
    for run in h.runs:
        set_chinese_font(run, font_size=16, bold=True)
    return h


def merge_documents():
    """合并所有Word文档"""
    base_path = r'F:\A-User\cliff\allin\软著申请材料'

    # 要合并的文档列表（按顺序）
    documents = [
        ('README.docx', '目录索引'),
        ('01_软件说明书.docx', '第一部分：软件说明书'),
        ('02_用户手册.docx', '第二部分：用户手册'),
        ('03_软件功能清单.docx', '第三部分：软件功能清单'),
        ('04_软件著作权申请表模板.docx', '第四部分：申请表模板'),
        ('05_源代码文档.docx', '第五部分：源代码文档'),
    ]

    # 创建新文档
    merged_doc = Document()

    # 设置页面
    section = merged_doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)

    # 添加封面
    cover_title = merged_doc.add_heading('微小型无人机智能风场测试评估系统', 0)
    for run in cover_title.runs:
        set_chinese_font(run, font_size=24, bold=True)

    cover_sub = merged_doc.add_heading('软件著作权申请材料', 0)
    for run in cover_sub.runs:
        set_chinese_font(run, font_size=20, bold=True)

    merged_doc.add_paragraph()

    # 封面信息表格
    info_table = merged_doc.add_table(rows=1, cols=2)
    info_table.style = 'Table Grid'

    cover_info = [
        ('软件名称', '微小型无人机智能风场测试评估系统'),
        ('软件简称', '风场测试评估系统'),
        ('版本号', 'V2.0.0'),
        ('开发完成日期', '2026年01月'),
        ('开发语言', 'Python'),
        ('代码行数', '约39,000行'),
    ]

    for row_data in cover_info:
        row = info_table.add_row()
        for i, value in enumerate(row_data):
            cell = row.cells[i]
            cell.text = str(value)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_chinese_font(run, font_size=12, bold=(i == 0))
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT

    merged_doc.add_paragraph()
    merged_doc.add_paragraph()

    # 目录
    merged_doc.add_heading('目录', level=1)

    for idx, (filename, section_title) in enumerate(documents, 1):
        p = merged_doc.add_paragraph(f'{idx}. {section_title}', style='List Number')

    merged_doc.add_page_break()

    # 合并各文档
    for idx, (filename, section_title) in enumerate(documents, 1):
        file_path = os.path.join(base_path, filename)

        if not os.path.exists(file_path):
            print(f'文件不存在: {filename}')
            continue

        print(f'正在合并: {filename}')

        # 添加分隔标题
        h = merged_doc.add_heading(section_title, level=1)
        for run in h.runs:
            set_chinese_font(run, font_size=16, bold=True)

        # 读取源文档
        try:
            source_doc = Document(file_path)

            # 复制内容（跳过第一个标题，避免重复）
            skip_first = True
            for element in source_doc.element.body:
                if element.tag.endswith('p'):  # 段落
                    # 查找对应的段落对象
                    for p in source_doc.paragraphs:
                        if p._element == element:
                            # 跳过第一个大标题
                            if skip_first and p.style and 'Heading' in str(p.style):
                                if p.style.name.startswith('Heading 1') and '微小型' in p.text or '软件著作权' in p.text:
                                    skip_first = False
                                    continue
                            copy_paragraph(p, merged_doc, keep_style=True)
                            break

                elif element.tag.endswith('tbl'):  # 表格
                    for table in source_doc.tables:
                        if table._element == element:
                            copy_table(table, merged_doc)
                            break

            # 添加分隔线
            merged_doc.add_paragraph('_' * 80)

        except Exception as e:
            print(f'合并 {filename} 时出错: {e}')

        # 添加分页符（除了最后一个文档）
        if idx < len(documents):
            merged_doc.add_page_break()

    # 保存合并后的文档
    output_path = os.path.join(base_path, '软著申请材料_完整版.docx')
    merged_doc.save(output_path)

    print()
    print(f'合并完成！')
    print(f'输出文件: {output_path}')

    return output_path


if __name__ == '__main__':
    merge_documents()
