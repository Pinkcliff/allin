# -*- coding: utf-8 -*-
"""
软件著作权申请材料 - Word文档生成器
将软著申请材料转换为Word格式
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# 设置中文字体
def set_chinese_font(run, font_name='宋体', font_size=12, bold=False):
    """设置中文字体"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    return run


def add_heading(doc, text, level=1):
    """添加标题"""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_chinese_font(run, font_size=16-level*2, bold=True)
    return h


def add_paragraph(doc, text, font_size=12, bold=False):
    """添加段落"""
    p = doc.add_paragraph(text)
    for run in p.runs:
        set_chinese_font(run, font_size=font_size, bold=bold)
    return p


def add_table_row(table, values, is_header=False):
    """添加表格行"""
    row = table.add_row()
    for i, value in enumerate(values):
        cell = row.cells[i]
        cell.text = str(value)
        # 设置单元格格式
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_chinese_font(run, font_size=11, bold=is_header)
            if is_header:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return row


def create_readme_doc():
    """生成README Word文档"""
    doc = Document()

    # 设置页面
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)

    # 标题
    add_heading(doc, '软件著作权申请材料', 0)
    add_heading(doc, '目录索引', 1)

    # 基本信息表格
    add_heading(doc, '软件基本信息', 2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    info = [
        ('项目', '内容'),
        ('软件名称', '微小型无人机智能风场测试评估系统'),
        ('软件简称', '风场测试评估系统'),
        ('版本号', 'V2.0.0'),
        ('开发完成日期', '2026年01月'),
        ('开发语言', 'Python'),
        ('代码行数', '约39,000行'),
    ]

    for row_data in info:
        add_table_row(table, row_data, is_header=(row_data[0] == '项目'))

    doc.add_paragraph()

    # 文件目录
    add_heading(doc, '文件目录', 2)
    files = [
        ('README.md', '目录索引'),
        ('01_软件说明书.docx', '软件功能详细说明'),
        ('02_用户手册.docx', '用户操作手册'),
        ('03_软件功能清单.docx', '功能模块清单'),
        ('04_软件著作权申请表模板.docx', '申请表填写模板'),
        ('05_源代码文档.docx', '源代码（前30页+后30页）'),
    ]

    for i, (file_name, desc) in enumerate(files, 1):
        add_paragraph(doc, f'{i}. {file_name}\t{desc}')

    doc.add_paragraph()

    # 核心功能概览
    add_heading(doc, '核心功能模块', 2)
    functions = [
        '全局监控仪表盘 - 系统资源、通讯状态、环境参数、日志中心监控',
        '风场编辑器 - 3D可视化编辑、风速区域、自定义表达式、模板库',
        '硬件控制 - 1600台风扇控制、Modbus/EtherCAT通讯',
        '传感器数据采集 - 1800+通道采集、Redis缓存、MongoDB同步',
        'CFD前处理 - 场景生成、网格生成、边界条件',
        '动捕数据可视化 - CSV加载、3D轨迹、播放控制、统计分析',
    ]

    for func in functions:
        add_paragraph(doc, '• ' + func)

    doc.add_paragraph()

    # 技术特性
    add_heading(doc, '技术特性', 2)
    tech_features = [
        '多线程架构 - UI与数据处理分离',
        '信号槽机制 - 模块间松耦合通信',
        '双层存储 - Redis + MongoDB',
        '实时3D渲染 - OpenGL高性能渲染',
        '表达式解析 - 动态数学表达式解析',
        '多协议支持 - TCP/IP、EtherCAT、Modbus',
    ]

    for feature in tech_features:
        add_paragraph(doc, '• ' + feature)

    doc.add_paragraph()

    # 性能指标
    add_heading(doc, '性能指标', 2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    metrics = [
        ('指标', '数值'),
        ('最大采集频率', '100Hz'),
        ('风扇响应时间', '<10ms'),
        ('3D渲染帧率', '>30fps'),
        ('支持传感器通道', '1800+'),
        ('数据存储容量', 'TB级'),
    ]

    for row_data in metrics:
        add_table_row(table, row_data, is_header=(row_data[0] == '指标'))

    # 保存
    output_path = r'F:\A-User\cliff\allin\软著申请材料\README.docx'
    doc.save(output_path)
    print(f'已生成: {output_path}')


def create_software_spec_doc():
    """生成软件说明书Word文档"""
    doc = Document()

    # 标题
    add_heading(doc, '微小型无人机智能风场测试评估系统', 0)
    add_heading(doc, '软件说明书', 1)

    # 一、软件基本信息
    add_heading(doc, '一、软件基本信息', 1)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    basic_info = [
        ('项目', '内容'),
        ('软件全称', '微小型无人机智能风场测试评估系统'),
        ('软件简称', '风场测试评估系统'),
        ('版本号', 'V2.0.0'),
        ('开发完成日期', '2026年01月'),
        ('首次发表日期', '2026年01月'),
        ('开发语言', 'Python 3.8+'),
        ('代码行数', '约39,000行'),
        ('运行环境', 'Windows 10/11, Linux'),
        ('软件性质', '原创'),
    ]

    for row_data in basic_info:
        add_table_row(table, row_data, is_header=(row_data[0] == '项目'))

    doc.add_paragraph()

    # 二、软件概述
    add_heading(doc, '二、软件概述', 1)

    add_heading(doc, '2.1 软件背景', 2)
    add_paragraph(doc, '本软件是一款面向微小型无人机风洞测试的专业综合评估系统，旨在解决无人机风场环境测试中的多设备协同控制、数据采集与分析、可视化监控等关键问题。系统集成了风扇阵列控制、环境参数监测、动作捕捉数据采集、CFD前处理等多个功能模块，为无人机风洞测试提供一体化解决方案。')

    add_heading(doc, '2.2 应用领域', 2)
    add_paragraph(doc, '• 无人机风洞试验测试')
    add_paragraph(doc, '• 空气动力学研究')
    add_paragraph(doc, '• 风场环境模拟与评估')
    add_paragraph(doc, '• 多传感器数据融合分析')
    add_paragraph(doc, '• 实验设备集成控制')

    add_heading(doc, '2.3 主要用途', 2)
    uses = [
        ('1. 风扇阵列控制', '控制1600台风扇生成各种风场模式'),
        ('2. 环境参数监测', '实时采集温度、湿度、风速、压力等传感器数据'),
        ('3. 动作捕捉集成', '接收和处理动捕系统数据，可视化无人机姿态轨迹'),
        ('4. 数据管理分析', '提供历史数据查询、统计分析、数据导出功能'),
        ('5. CFD前处理', '为计算流体力学仿真生成网格和场景文件'),
        ('6. 系统监控', '统一监控各子系统运行状态'),
    ]

    for title, desc in uses:
        add_paragraph(doc, f'{title}：{desc}')

    doc.add_paragraph()

    # 三、功能说明
    add_heading(doc, '三、功能说明', 1)

    add_heading(doc, '3.1 系统架构', 2)
    add_paragraph(doc, '本软件采用模块化设计，主要包含以下子系统：')
    add_paragraph(doc, '• 融合平台主入口')
    add_paragraph(doc, '• 全局监控仪表盘')
    add_paragraph(doc, '• 风场编辑器')
    add_paragraph(doc, '• 硬件控制模块')
    add_paragraph(doc, '• 传感器数据采集模块')
    add_paragraph(doc, '• CFD前处理模块')
    add_paragraph(doc, '• 动捕数据可视化模块')
    add_paragraph(doc, '• 核心数据层（Redis/MongoDB）')

    add_heading(doc, '3.2 核心功能模块', 2)

    add_heading(doc, '3.2.1 全局监控仪表盘', 3)
    dashboard_features = [
        '系统监控Dock：显示CPU、内存、网络等系统资源使用情况',
        '通讯状态Dock：监控各通讯协议的连接状态（TCP/IP、EtherCAT、Modbus等）',
        '环境监测Dock：实时显示温度、湿度、压力、风速等环境参数',
        '日志中心Dock：系统运行日志的实时显示和过滤',
        '动捕数据Dock：显示动作捕捉系统的实时数据',
        '风机状态Dock：1600台风扇的PWM值监控和热力图显示',
    ]
    for feature in dashboard_features:
        add_paragraph(doc, '• ' + feature)

    add_heading(doc, '3.2.2 风场编辑器', 3)
    editor_features = [
        '3D可视化编辑：在3D场景中布置风速区域和目标点',
        '表达式解析：支持自定义数学表达式定义风速分布',
        '模板库管理：预设多种常用风场模式（层流、湍流、梯度风等）',
        '时间轴控制：支持风场随时间变化的动画编辑',
        '实时预览：实时查看风场矢量和风速分布',
    ]
    for feature in editor_features:
        add_paragraph(doc, '• ' + feature)

    add_heading(doc, '3.2.3 硬件控制模块', 3)
    add_paragraph(doc, '实现对各类硬件设备的统一控制：')

    doc.add_paragraph('• 风扇阵列控制：')
    add_paragraph(doc, '  - 支持1600台风扇（40×40阵列）的独立PWM控制')
    add_paragraph(doc, '  - 提供批量控制、区域控制、梯度控制等多种控制模式')
    add_paragraph(doc, '  - 支持CSV格式控制序列导入和回放')

    doc.add_paragraph('• Modbus设备控制：')
    add_paragraph(doc, '  - 支持俯仰伺服、造雨装置、喷雾装置等设备')
    add_paragraph(doc, '  - 标准Modbus RTU/TCP协议')

    doc.add_paragraph('• EtherCAT设备控制：')
    add_paragraph(doc, '  - 支持电驱、风速传感器、温度传感器等')
    add_paragraph(doc, '  - 微秒级实时控制周期')

    add_heading(doc, '3.2.4 传感器数据采集', 3)
    add_paragraph(doc, '提供完整的传感器数据采集和管理功能：')

    sensor_types = [
        ('1600路', '风扇PWM值'),
        ('100路', '温度传感器（-20~80℃）'),
        ('100路', '风速传感器（0~30m/s）'),
        ('4路', '温湿度传感器'),
        ('1路', '大气压力传感器'),
    ]
    for count, desc in sensor_types:
        add_paragraph(doc, f'• {count} {desc}')

    doc.add_paragraph()
    add_paragraph(doc, '数据存储与管理：')
    add_paragraph(doc, '• Redis实时数据缓存')
    add_paragraph(doc, '• MongoDB历史数据归档')
    add_paragraph(doc, '• CSV格式导出')
    add_paragraph(doc, '• 历史数据查询和图表显示')
    add_paragraph(doc, '• 统计分析（最大值、最小值、平均值、标准差等）')

    add_heading(doc, '3.2.5 CFD前处理模块', 3)
    cfd_features = [
        '场景生成：根据风场配置生成CFD仿真场景文件',
        '网格生成：自动生成计算域网格',
        '边界条件设置：自动设置入口、出口、壁面等边界条件',
        '风扇模型映射：将风扇阵列映射到CFD边界条件',
    ]
    for feature in cfd_features:
        add_paragraph(doc, '• ' + feature)

    add_heading(doc, '3.2.6 动捕数据可视化', 3)
    motion_features = [
        'CSV数据加载：支持LuMo动捕系统导出的CSV格式',
        '3D轨迹显示：标记点的3D空间轨迹可视化',
        '播放控制：支持播放、暂停、停止、跳转功能',
        '统计分析：每个标记点的位置统计分析',
        '数据导出：支持导出处理后的数据',
    ]
    for feature in motion_features:
        add_paragraph(doc, '• ' + feature)

    doc.add_paragraph()

    # 四、技术特点
    add_heading(doc, '四、技术特点', 1)

    add_heading(doc, '4.1 技术架构', 2)
    tech_stack = [
        ('开发语言', 'Python 3.8+'),
        ('GUI框架', 'PySide6 (Qt6)'),
        ('数据可视化', 'Matplotlib, PyQtGraph'),
        ('数据存储', 'Redis, MongoDB'),
        ('通讯协议', 'Socket, Modbus-tk, EtherCAT'),
    ]
    for item, value in tech_stack:
        add_paragraph(doc, f'• {item}：{value}')

    doc.add_paragraph()

    add_heading(doc, '4.2 关键技术', 2)
    key_tech = [
        '多线程架构：采用工作线程模式，保证UI响应流畅',
        '信号槽机制：基于Qt信号槽实现模块间松耦合通信',
        '数据缓存策略：Redis实时缓存 + MongoDB持久化的双层存储',
        '表达式解析引擎：支持复杂数学表达式的动态解析和计算',
        '3D可视化技术：基于OpenGL的高性能3D渲染',
    ]
    for tech in key_tech:
        add_paragraph(doc, '• ' + tech)

    doc.add_paragraph()

    add_heading(doc, '4.3 性能指标', 2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    perf_data = [
        ('指标', '数值'),
        ('数据采集频率', '最高100Hz'),
        ('风扇响应时间', '<10ms'),
        ('3D渲染帧率', '>30fps'),
        ('支持传感器通道数', '>1800路'),
        ('数据存储容量', 'TB级'),
    ]

    for row_data in perf_data:
        add_table_row(table, row_data, is_header=(row_data[0] == '指标'))

    doc.add_paragraph()

    add_heading(doc, '4.4 创新点', 2)
    innovations = [
        '一体化融合平台：首次将风场控制、数据采集、动捕、CFD等功能集成于统一平台',
        '大规模阵列控制：实现1600台风扇的独立精确控制',
        '自定义风场表达式：支持用户通过数学表达式定义任意风场分布',
        '多源数据融合：整合传感器、动捕、仿真等多源数据',
        '实时3D可视化：提供直观的3D风场可视化编辑界面',
    ]
    for i, innovation in enumerate(innovations, 1):
        add_paragraph(doc, f'{i}. {innovation}')

    doc.add_paragraph()

    # 五、系统要求
    add_heading(doc, '五、系统要求', 1)

    add_heading(doc, '5.1 硬件要求', 2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    hw_req = [
        ('配置', '最低要求', '推荐配置'),
        ('CPU', 'Intel i5或同等', 'Intel i7或更高'),
        ('内存', '8GB', '16GB以上'),
        ('硬盘', '100GB可用空间', 'SSD 500GB以上'),
        ('网络', '千兆网卡', '万兆网卡'),
        ('显示器', '1920×1080', '2560×1440以上'),
    ]

    for row_data in hw_req:
        add_table_row(table, row_data, is_header=(row_data[0] == '配置'))

    doc.add_paragraph()

    add_heading(doc, '5.2 软件要求', 2)
    add_paragraph(doc, '• 操作系统：Windows 10/11 (64位) 或 Linux (Ubuntu 20.04+)')
    add_paragraph(doc, '• Python版本：3.8 或更高')
    add_paragraph(doc, '• 数据库：Redis 6.0+, MongoDB 4.4+ (可选)')

    doc.add_paragraph()
    add_paragraph(doc, '依赖库：')
    libs = ['PySide6', 'redis', 'pymongo', 'matplotlib', 'numpy', 'pandas', 'pyqtgraph']
    for lib in libs:
        add_paragraph(doc, f'  - {lib}')

    doc.add_paragraph()

    # 六、版权声明
    add_heading(doc, '六、版权声明', 1)
    add_paragraph(doc, '本软件《微小型无人机智能风场测试评估系统》V2.0.0 的著作权归开发团队所有。未经著作权人书面许可，任何单位或个人不得以任何形式复制、修改、传播、出租、出售本软件或利用本软件从事商业活动。')
    doc.add_paragraph()
    add_paragraph(doc, '本软件受《中华人民共和国著作权法》及国际著作权条约的保护。')

    doc.add_paragraph()
    add_paragraph(doc, '本说明书最后更新日期：2026年01月')

    # 保存
    output_path = r'F:\A-User\cliff\allin\软著申请材料\01_软件说明书.docx'
    doc.save(output_path)
    print(f'已生成: {output_path}')


def create_function_list_doc():
    """生成软件功能清单Word文档"""
    doc = Document()

    # 标题
    add_heading(doc, '微小型无人机智能风场测试评估系统', 0)
    add_heading(doc, '软件功能清单', 1)

    # 一、软件基本信息
    add_heading(doc, '一、软件基本信息', 1)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    basic_info = [
        ('项目', '内容'),
        ('软件全称', '微小型无人机智能风场测试评估系统'),
        ('软件简称', '风场测试评估系统'),
        ('版本号', 'V2.0.0'),
        ('开发完成日期', '2026年01月'),
        ('开发语言', 'Python'),
        ('代码行数', '约39,000行'),
    ]

    for row_data in basic_info:
        add_table_row(table, row_data, is_header=(row_data[0] == '项目'))

    doc.add_paragraph()

    # 二、功能模块清单
    add_heading(doc, '二、功能模块清单', 1)

    modules = [
        ('2.1 融合平台主入口模块', [
            ('1', '启动器界面', '提供模块选择入口'),
            ('2', '全局仪表盘启动', '启动完整监控仪表盘'),
            ('3', '融合系统启动', '启动仪表盘+传感器数据'),
            ('4', '传感器采集启动', '启动独立数据采集系统'),
        ]),
        ('2.2 全局监控仪表盘模块', [
            ('5', '系统资源监控', 'CPU、内存、网络使用率实时显示'),
            ('6', '通讯状态监控', 'TCP/IP、EtherCAT、Modbus连接状态'),
            ('7', '环境参数监控', '温度、湿度、风速、压力实时显示'),
            ('8', '日志中心', '系统日志实时显示、过滤、搜索'),
            ('9', '动捕数据监控', '动作捕捉数据实时显示'),
            ('10', '风机状态监控', '1600台风扇PWM值监控和热力图'),
            ('11', 'Dock窗口管理', '可拖拽、可停靠的窗口布局'),
            ('12', '主题管理', '深色/浅色主题切换'),
        ]),
        ('2.3 风场编辑器模块', [
            ('13', '3D场景可视化', '基于OpenGL的3D风场显示'),
            ('14', '风速区域添加', '在3D场景中添加矩形/圆形区域'),
            ('15', '区域属性编辑', '位置、尺寸、风速表达式编辑'),
            ('16', '自定义表达式', '支持复杂数学表达式定义风速'),
            ('17', '表达式解析器', '动态解析和计算数学表达式'),
            ('18', '风场模板库', '预设层流、湍流、梯度风等模板'),
            ('19', '时间轴编辑', '支持风场随时间变化的动画'),
            ('20', '关键帧管理', '添加、编辑、删除关键帧'),
            ('21', '插值计算', '关键帧间自动插值'),
            ('22', '风场预览', '实时预览风场矢量和分布'),
            ('23', '项目保存加载', '保存和加载风场配置文件'),
            ('24', 'CSV导出', '导出风速分布数据'),
        ]),
        ('2.4 硬件控制模块', [
            ('25', 'Modbus通讯', '支持Modbus RTU/TCP协议'),
            ('26', 'EtherCAT通讯', '支持EtherCAT实时通讯'),
            ('27', 'TCP/IP通讯', '支持标准TCP/IP通讯'),
            ('28', '风扇阵列控制', '控制1600台风扇(40×40阵列)'),
            ('29', '单风扇控制', '单个风扇PWM值设置'),
            ('30', '批量风扇控制', '批量设置多台风扇PWM值'),
            ('31', '区域风扇控制', '按区域设置风扇PWM值'),
            ('32', '梯度风扇控制', '设置PWM梯度分布'),
            ('33', 'CSV控制序列', '加载和回放CSV控制文件'),
            ('34', '俯仰伺服控制', 'Modbus伺服电机控制'),
            ('35', '造雨装置控制', 'Modbus造雨设备控制'),
            ('36', '喷雾装置控制', 'Modbus喷雾设备控制'),
            ('37', '设备状态监测', '实时监测设备运行状态'),
            ('38', '控制日志记录', '记录控制操作历史'),
            ('39', '错误处理', '设备故障检测和处理'),
        ]),
        ('2.5 传感器数据采集模块', [
            ('40', '多通道采集', '支持1800+传感器通道'),
            ('41', '风扇PWM采集', '1600路风扇PWM值采集'),
            ('42', '温度采集', '100路温度传感器采集'),
            ('43', '风速采集', '100路风速传感器采集'),
            ('44', '温湿度采集', '4路温湿度传感器采集'),
            ('45', '压力采集', '大气压力传感器采集'),
            ('46', '采样率配置', '可配置采样频率'),
            ('47', '定时采集', '支持定时自动采集'),
            ('48', '采集进度显示', '实时显示采集进度'),
            ('49', 'Redis缓存', '采集数据实时缓存到Redis'),
            ('50', 'MongoDB同步', '数据同步到MongoDB持久化'),
            ('51', '历史数据查询', '查询历史采集记录'),
            ('52', '数据统计', '最大值、最小值、平均值等统计'),
            ('53', '数据图表', '数据趋势图表显示'),
            ('54', 'CSV导出', '导出为CSV格式文件'),
        ]),
        ('2.6 CFD前处理模块', [
            ('55', '场景生成', '根据风场配置生成CFD场景'),
            ('56', '网格生成', '自动生成计算域网格'),
            ('57', '网格密度配置', '可配置网格分辨率'),
            ('58', '边界条件设置', '自动设置入口/出口边界'),
            ('59', '风扇模型映射', '风扇阵列映射到CFD边界'),
            ('60', 'VTM文件导出', '导出VTK多块数据集'),
            ('61', '场景预览', '预览生成的CFD场景'),
        ]),
        ('2.7 动捕数据可视化模块', [
            ('62', 'CSV数据加载', '加载LuMo动捕系统CSV文件'),
            ('63', '3D轨迹显示', '标记点3D轨迹可视化'),
            ('64', '播放控制', '播放/暂停/停止控制'),
            ('65', '跳转控制', '滑块跳转到指定帧'),
            ('66', '速度控制', '可调节播放速度'),
            ('67', '标记点统计', '每个标记点位置统计'),
            ('68', '直方图分析', '位置分布直方图'),
            ('69', '数据表格', '帧数据表格显示'),
            ('70', '数据导出', '导出处理后数据'),
        ]),
        ('2.8 数据存储模块', [
            ('71', 'Redis连接', 'Redis数据库连接管理'),
            ('72', '数据写入', '实时数据写入Redis'),
            ('73', '数据读取', 'Redis数据读取'),
            ('74', '数据订阅', 'Redis数据订阅/发布'),
            ('75', 'MongoDB连接', 'MongoDB连接管理'),
            ('76', '批量同步', '批量数据同步到MongoDB'),
            ('77', '进度显示', '同步进度显示'),
            ('78', '索引创建', '自动创建数据索引'),
            ('79', '数据库管理', '数据库连接配置'),
        ]),
        ('2.9 工具模块', [
            ('80', '配置管理', '统一配置文件管理'),
            ('81', '日志系统', '分级日志记录'),
            ('82', '信号通信', 'Qt信号槽通信'),
            ('83', '线程管理', '工作线程管理'),
            ('84', '错误处理', '统一错误处理机制'),
            ('85', '主题样式', 'UI主题管理'),
            ('86', '自定义控件', '自定义UI控件库'),
        ]),
    ]

    for module_title, functions in modules:
        add_heading(doc, module_title, 2)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'

        # 表头
        add_table_row(table, ['序号', '功能名称', '功能描述'], is_header=True)

        # 功能行
        for func_data in functions:
            add_table_row(table, func_data)

        doc.add_paragraph()

    # 三、技术特性清单
    add_heading(doc, '三、技术特性清单', 1)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    tech_features = [
        ('序号', '技术特性'),
        ('1', '多线程架构 - UI与数据处理分离'),
        ('2', '信号槽机制 - 模块间松耦合通信'),
        ('3', '双层存储 - Redis+MongoDB双层存储'),
        ('4', '实时3D渲染 - OpenGL高性能渲染'),
        ('5', '表达式解析 - 动态数学表达式解析'),
        ('6', '多协议支持 - TCP/IP/EtherCAT/Modbus'),
        ('7', '模块化设计 - 高内聚低耦合架构'),
        ('8', '插件式扩展 - 支持功能模块扩展'),
    ]

    for row_data in tech_features:
        add_table_row(table, row_data, is_header=(row_data[0] == '序号'))

    doc.add_paragraph()

    # 四、性能指标清单
    add_heading(doc, '四、性能指标清单', 1)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    perf_data = [
        ('指标项', '指标值'),
        ('最大采集频率', '100Hz'),
        ('风扇响应时间', '<10ms'),
        ('3D渲染帧率', '>30fps'),
        ('支持传感器通道', '1800+'),
        ('数据存储容量', 'TB级'),
        ('并发控制设备', '100+'),
        ('网络通讯延迟', '<5ms'),
    ]

    for row_data in perf_data:
        add_table_row(table, row_data, is_header=(row_data[0] == '指标项'))

    doc.add_paragraph()

    # 五、统计汇总
    add_heading(doc, '五、统计汇总', 1)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    summary = [
        ('类别', '数量'),
        ('功能模块数', '9个'),
        ('功能点数', '86个'),
        ('Python源文件', '100+个'),
        ('代码总行数', '约39,000行'),
        ('支持的通讯协议', '4种'),
        ('支持的传感器类型', '5种'),
    ]

    for row_data in summary:
        add_table_row(table, row_data, is_header=(row_data[0] == '类别'))

    doc.add_paragraph()
    add_paragraph(doc, '功能清单文档 版本 2.0.0')
    add_paragraph(doc, '生成日期：2026年01月')

    # 保存
    output_path = r'F:\A-User\cliff\allin\软著申请材料\03_软件功能清单.docx'
    doc.save(output_path)
    print(f'已生成: {output_path}')


def create_application_form_doc():
    """生成申请表模板Word文档"""
    doc = Document()

    # 标题
    add_heading(doc, '计算机软件著作权登记申请表', 0)

    # 一、软件基本信息
    add_heading(doc, '一、软件基本信息', 1)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    sw_info = [
        ('申请项目', '填写内容'),
        ('软件全称', '微小型无人机智能风场测试评估系统'),
        ('软件简称', '风场测试评估系统（可选填）'),
        ('版本号', 'V2.0.0'),
        ('分类号', '(按系统自动生成)'),
        ('开发完成日期', '2026年01月__日'),
        ('首次发表日期', '2026年01月__日'),
    ]

    for row_data in sw_info:
        add_table_row(table, row_data, is_header=(row_data[0] == '申请项目'))

    doc.add_paragraph()

    # 二、著作权人信息
    add_heading(doc, '二、著作权人信息', 1)

    add_heading(doc, '2.1 著作权人（申请人）', 2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    owner_info = [
        ('申请项目', '填写内容'),
        ('姓名或名称', '【请填写】'),
        ('性质', '□ 法人 □ 非法人组织 □ 自然人 □ 其他'),
        ('证件类型', '□ 营业执照 □ 事业单位法人证书 □ 身份证 □ 其他'),
        ('证件号码', '【请填写】'),
        ('国籍', '中国'),
        ('省份/城市', '【请填写】'),
        ('行政区代码', '【请填写】'),
        ('地址', '【请填写详细地址】'),
        ('邮政编码', '【请填写】'),
        ('联系人', '【请填写】'),
        ('联系电话', '【请填写】'),
        ('电子邮箱', '【请填写】'),
    ]

    for row_data in owner_info:
        add_table_row(table, row_data, is_header=(row_data[0] == '申请项目'))

    doc.add_paragraph()

    # 三、软件描述
    add_heading(doc, '三、软件描述', 1)

    add_heading(doc, '3.1 功能简介', 2)
    add_paragraph(doc, '本软件是一款面向微小型无人机风洞测试的专业综合评估系统，集成了')
    add_paragraph(doc, '风扇阵列控制、环境参数监测、动作捕捉数据采集、CFD前处理等多个')
    add_paragraph(doc, '功能模块。系统支持1600台风扇的独立PWM控制，1800+传感器通道的')
    add_paragraph(doc, '数据采集，提供3D可视化风场编辑界面，支持多种通讯协议（TCP/IP、')
    add_paragraph(doc, 'EtherCAT、Modbus），实现测试设备一体化控制和数据管理。')

    doc.add_paragraph()

    add_heading(doc, '3.2 主要功能', 2)
    main_funcs = [
        '1. 全局监控仪表盘：统一监控各子系统运行状态',
        '2. 风场编辑器：可视化配置风场参数，支持自定义表达式',
        '3. 硬件控制模块：风扇阵列、伺服电机等设备控制',
        '4. 传感器数据采集：温度、风速、湿度等多通道数据采集',
        '5. CFD前处理：生成计算流体力学仿真场景文件',
        '6. 动捕数据可视化：动作捕捉数据的3D可视化和分析',
    ]
    for func in main_funcs:
        add_paragraph(doc, func)

    doc.add_paragraph()

    add_heading(doc, '3.3 技术特点', 2)
    tech_points = [
        '1. 模块化设计，支持灵活配置和扩展',
        '2. 多线程架构，保证系统响应性能',
        '3. 双层数据存储（Redis+MongoDB），支持大数据量',
        '4. 实时3D可视化，直观显示风场分布',
        '5. 支持多种工业通讯协议',
        '6. 自定义表达式引擎，灵活定义风场参数',
    ]
    for point in tech_points:
        add_paragraph(doc, point)

    doc.add_paragraph()

    # 四、开发信息
    add_heading(doc, '四、软件开发信息', 1)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    dev_info = [
        ('开发方式', '□ 独立开发 √ □ 合作开发 □ 委托开发 □ 下达任务开发'),
        ('开发完成日期', '2026年01月__日'),
        ('首次发表日期', '2026年01月__日'),
        ('开发环境', 'Python 3.8+ / PySide6 / Windows 10/11'),
        ('硬件环境', 'CPU i5+/ 内存8GB+/ 硬盘100GB+'),
        ('编程语言', 'Python'),
        ('代码行数', '约39,000行'),
    ]

    for row_data in dev_info:
        add_table_row(table, row_data, is_header=(row_data[0] == '开发方式'))

    doc.add_paragraph()

    # 五、权利状况
    add_heading(doc, '五、权利状况', 1)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    rights_info = [
        ('权利取得方式', '□ 原始取得 √ □ 继受取得'),
        ('权利范围', '全部权利 √'),
    ]

    for row_data in rights_info:
        add_table_row(table, row_data)

    doc.add_paragraph()

    # 六、申请人承诺
    add_heading(doc, '六、申请人承诺', 1)
    add_paragraph(doc, '本人/本单位保证：', bold=True)
    promises = [
        '1. 所填内容及提交的申请文件、材料真实、准确；',
        '2. 本软件不含有任何违反法律法规的内容；',
        '3. 本软件未侵犯任何第三方的著作权或其他合法权益；',
        '4. 如有违反，本人/本单位愿意承担一切法律责任。',
    ]
    for promise in promises:
        add_paragraph(doc, promise)

    doc.add_paragraph()
    doc.add_paragraph()
    add_paragraph(doc, f'申请人签字/盖章：____________________', font_size=12)
    add_paragraph(doc, f'日期：2026年01月__日', font_size=12)

    doc.add_paragraph()
    doc.add_paragraph()
    add_paragraph(doc, '申请表模板 版本 2.0.0')
    add_paragraph(doc, '填写日期：2026年01月')

    # 保存
    output_path = r'F:\A-User\cliff\allin\软著申请材料\04_软件著作权申请表模板.docx'
    doc.save(output_path)
    print(f'已生成: {output_path}')


def create_user_manual_doc():
    """生成用户手册Word文档"""
    doc = Document()

    # 标题
    add_heading(doc, '微小型无人机智能风场测试评估系统', 0)
    add_heading(doc, '用户手册', 1)

    # 一、系统简介
    add_heading(doc, '一、系统简介', 1)

    add_heading(doc, '1.1 系统概述', 2)
    add_paragraph(doc, '本手册面向《微小型无人机智能风场测试评估系统》V2.0.0 的用户，介绍系统的安装、配置和操作方法。')

    add_heading(doc, '1.2 主要功能', 2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    funcs = [
        ('功能模块', '描述'),
        ('全局监控仪表盘', '统一监控各子系统状态'),
        ('风场编辑器', '可视化配置风场参数'),
        ('硬件控制', '控制风扇阵列等设备'),
        ('传感器数据采集', '采集和存储传感器数据'),
        ('CFD前处理', '生成仿真前处理文件'),
        ('动捕数据可视化', '处理动作捕捉数据'),
    ]

    for row_data in funcs:
        add_table_row(table, row_data, is_header=(row_data[0] == '功能模块'))

    doc.add_paragraph()

    # 二、安装与配置
    add_heading(doc, '二、安装与配置', 1)

    add_heading(doc, '2.1 系统要求', 2)
    add_paragraph(doc, '硬件要求：')
    add_paragraph(doc, '• CPU: Intel i5 或同等性能处理器')
    add_paragraph(doc, '• 内存: 8GB 以上（推荐 16GB）')
    add_paragraph(doc, '• 硬盘: 100GB 可用空间')
    add_paragraph(doc, '• 网络: 千兆以太网')

    doc.add_paragraph()
    add_paragraph(doc, '软件要求：')
    add_paragraph(doc, '• 操作系统: Windows 10/11 (64位) 或 Linux')
    add_paragraph(doc, '• Python: 3.8 或更高版本')
    add_paragraph(doc, '• 数据库: Redis 6.0+, MongoDB 4.4+ (可选)')

    doc.add_paragraph()

    add_heading(doc, '2.2 安装步骤', 2)
    add_paragraph(doc, '1. 安装Python环境')
    add_paragraph(doc, '   下载并安装 Python 3.8+，确保 pip 可用')

    doc.add_paragraph()
    add_paragraph(doc, '2. 安装依赖库')
    add_paragraph(doc, '   pip install PySide6 redis pymongo matplotlib numpy pandas pyqtgraph')

    doc.add_paragraph()
    add_paragraph(doc, '3. 配置系统参数')
    add_paragraph(doc, '   编辑 config.py 文件，设置通讯参数、数据库连接等')

    doc.add_paragraph()

    # 三、快速开始
    add_heading(doc, '三、快速开始', 1)

    add_heading(doc, '3.1 启动系统', 2)
    add_paragraph(doc, '双击运行 main.py 或在命令行执行：')
    add_paragraph(doc, 'python main.py')

    doc.add_paragraph()
    add_paragraph(doc, '系统将显示启动器界面，提供三个选项：')

    launch_options = [
        ('1', '全局监控仪表盘', '完整的监控界面'),
        ('2', '融合系统', '仪表盘 + 传感器数据'),
        ('3', '传感器数据采集', '独立的数据采集界面'),
    ]
    for num, name, desc in launch_options:
        add_paragraph(doc, f'{num}. {name} - {desc}')

    doc.add_paragraph()

    add_heading(doc, '3.2 界面介绍', 2)
    add_paragraph(doc, '主界面布局：')
    add_paragraph(doc, '• 顶部：菜单栏和工具栏')
    add_paragraph(doc, '• 中间：主显示区域 + Dock窗口区域')
    add_paragraph(doc, '• 底部：状态栏')

    doc.add_paragraph()

    # 四、功能模块使用说明
    add_heading(doc, '四、功能模块使用说明', 1)

    add_heading(doc, '4.1 全局监控仪表盘', 2)
    dashboard_usage = [
        ('系统监控', '显示CPU、内存、网络使用情况'),
        ('通讯状态', '显示各通讯协议连接状态'),
        ('环境监测', '显示温度、湿度、风速、压力等参数'),
        ('日志中心', '查看系统运行日志'),
        ('动捕数据', '显示动作捕捉实时数据'),
        ('风机状态', '显示1600台风扇PWM值'),
    ]
    for title, desc in dashboard_usage:
        add_paragraph(doc, f'• {title}：{desc}')

    doc.add_paragraph()

    add_heading(doc, '4.2 风场编辑器', 2)
    add_paragraph(doc, '新建风场配置：')
    add_paragraph(doc, '1. 点击"新建"按钮')
    add_paragraph(doc, '2. 设置风场参数（名称、描述）')
    add_paragraph(doc, '3. 在3D视图中添加风速区域')

    doc.add_paragraph()
    add_paragraph(doc, '添加风速区域：')
    add_paragraph(doc, '1. 选择"添加区域"工具')
    add_paragraph(doc, '2. 在3D场景中点击确定区域位置')
    add_paragraph(doc, '3. 设置区域参数（位置、尺寸、风速表达式）')

    doc.add_paragraph()
    add_paragraph(doc, '保存和导出：')
    add_paragraph(doc, '• 保存项目：保存为 .wf 格式')
    add_paragraph(doc, '• 导出CSV：导出风速分布数据')
    add_paragraph(doc, '• 导出控制序列：导出到硬件控制系统')

    doc.add_paragraph()

    add_heading(doc, '4.3 硬件控制', 2)
    add_paragraph(doc, '连接设备：')
    add_paragraph(doc, '1. 打开硬件控制面板')
    add_paragraph(doc, '2. 配置通讯参数')
    add_paragraph(doc, '3. 点击"连接"按钮')

    doc.add_paragraph()
    add_paragraph(doc, '风扇控制：')
    add_paragraph(doc, '• 单个控制：选择风扇，设置PWM值，点击应用')
    add_paragraph(doc, '• 区域控制：选择区域，设置PWM值，点击应用')
    add_paragraph(doc, '• 梯度控制：设置起终点PWM值，选择方向，点击应用')

    doc.add_paragraph()

    add_heading(doc, '4.4 传感器数据采集', 2)
    add_paragraph(doc, '创建采集任务：')
    add_paragraph(doc, '1. 打开传感器数据采集窗口')
    add_paragraph(doc, '2. 设置采集参数（采样率、时长）')
    add_paragraph(doc, '3. 点击"开始采集"')

    doc.add_paragraph()
    add_paragraph(doc, '查看历史数据：')
    add_paragraph(doc, '1. 切换到"数据查看"标签页')
    add_paragraph(doc, '2. 选择要查看的采集记录')
    add_paragraph(doc, '3. 查看数据详情和统计图表')

    doc.add_paragraph()

    # 五、常见问题
    add_heading(doc, '五、常见问题', 1)

    add_paragraph(doc, '问题1：系统无法启动', font_size=12, bold=True)
    add_paragraph(doc, '解决方案：')
    add_paragraph(doc, '1. 检查Python是否正确安装')
    add_paragraph(doc, '2. 检查依赖库是否安装完整')
    add_paragraph(doc, '3. 在命令行运行查看错误信息')

    doc.add_paragraph()

    add_paragraph(doc, '问题2：无法连接硬件设备', font_size=12, bold=True)
    add_paragraph(doc, '解决方案：')
    add_paragraph(doc, '1. 检查网络连接')
    add_paragraph(doc, '2. 确认设备IP地址配置正确')
    add_paragraph(doc, '3. 检查防火墙设置')
    add_paragraph(doc, '4. 确认设备已上电')

    doc.add_paragraph()

    add_paragraph(doc, '问题3：数据采集失败', font_size=12, bold=True)
    add_paragraph(doc, '解决方案：')
    add_paragraph(doc, '1. 检查Redis服务是否启动')
    add_paragraph(doc, '2. 确认传感器连接正常')
    add_paragraph(doc, '3. 检查磁盘空间是否充足')

    doc.add_paragraph()

    # 六、附录
    add_heading(doc, '六、附录', 1)

    add_heading(doc, '6.1 快捷键', 2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    shortcuts = [
        ('快捷键', '功能'),
        ('Ctrl+N', '新建项目'),
        ('Ctrl+O', '打开项目'),
        ('Ctrl+S', '保存项目'),
        ('Space', '播放/暂停'),
        ('ESC', '停止'),
        ('F1', '帮助'),
    ]

    for row_data in shortcuts:
        add_table_row(table, row_data, is_header=(row_data[0] == '快捷键'))

    doc.add_paragraph()

    add_heading(doc, '6.2 文件格式说明', 2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    formats = [
        ('扩展名', '描述'),
        ('.wf', '风场配置文件'),
        ('.csv', '逗号分隔值数据文件'),
        ('.vtm', 'VTK多块数据集'),
        ('.log', '日志文件'),
    ]

    for row_data in formats:
        add_table_row(table, row_data, is_header=(row_data[0] == '扩展名'))

    doc.add_paragraph()
    add_paragraph(doc, '用户手册 版本 2.0.0')
    add_paragraph(doc, '最后更新：2026年01月')

    # 保存
    output_path = r'F:\A-User\cliff\allin\软著申请材料\02_用户手册.docx'
    doc.save(output_path)
    print(f'已生成: {output_path}')


def create_source_code_doc():
    """生成源代码Word文档"""
    # 读取已生成的源代码文档
    source_file = r'F:\A-User\cliff\allin\软著申请材料\源代码文档.txt'

    if not os.path.exists(source_file):
        print(f'源代码文档不存在，请先运行生成源代码文档脚本')
        return

    with open(source_file, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()

    # 标题
    add_heading(doc, '微小型无人机智能风场测试评估系统', 0)
    add_heading(doc, '源代码文档', 1)

    add_paragraph(doc, '软件名称：微小型无人机智能风场测试评估系统')
    add_paragraph(doc, '版本号：V2.0.0')
    add_paragraph(doc, '文档类型：源代码文档')
    add_paragraph(doc, '生成日期：2026年01月')

    doc.add_paragraph()

    # 按行分割并添加到文档
    lines = content.split('\n')

    # 跳过前面的标题部分，从实际代码开始
    in_code_section = False
    line_count = 0
    page_line_count = 0
    current_file = None

    for line in lines:
        # 检测是否进入代码部分
        if '前30页源代码' in line or '后30页源代码' in line:
            in_code_section = True
            doc.add_paragraph()
            add_paragraph(doc, line.replace('=', ''), bold=True)
            doc.add_paragraph()
            continue

        if '===' in line and '源代码文档结束' in line:
            break

        # 检测文件头
        if '文件：' in line and 'all' in line:
            if current_file != line:
                current_file = line
                doc.add_page_break()
                add_paragraph(doc, line.strip('-').strip(), bold=True)
                add_paragraph(doc, f'[第 {page_count := page_count + 1} 页]' if 'page_count' not in locals() else '')
            continue

        if '[第' in line and '页' in line:
            add_paragraph(doc, line.strip())
            continue

        # 添加代码行
        if in_code_section and line.strip():
            # 添加带行号的代码行
            p = doc.add_paragraph()
            run = p.add_run(line)
            set_chinese_font(run, font_name='Consolas', font_size=10)
            line_count += 1
            page_line_count += 1

            # 每50行分页（软著要求）
            if page_line_count >= 50:
                doc.add_page_break()
                page_line_count = 0

    doc.add_paragraph()
    add_paragraph(doc, f'源代码文档结束，共 {line_count} 行')

    # 保存
    output_path = r'F:\A-User\cliff\allin\软著申请材料\05_源代码文档.docx'
    doc.save(output_path)
    print(f'已生成: {output_path}')


def main():
    """主函数"""
    print('正在生成Word文档...')
    print()

    create_readme_doc()
    create_software_spec_doc()
    create_function_list_doc()
    create_application_form_doc()
    create_user_manual_doc()
    create_source_code_doc()

    print()
    print('所有Word文档生成完成！')
    print()
    print('生成的文件：')
    print('  - README.docx')
    print('  - 01_软件说明书.docx')
    print('  - 02_用户手册.docx')
    print('  - 03_软件功能清单.docx')
    print('  - 04_软件著作权申请表模板.docx')
    print('  - 05_源代码文档.docx')


if __name__ == '__main__':
    # 定义全局变量
    page_count = 0

    main()
