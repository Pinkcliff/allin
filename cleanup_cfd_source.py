# -*- coding: utf-8 -*-
"""清理源代码中的CFD模块内容"""

from docx import Document

doc = Document(r'F:\A-User\cliff\allin\软著申请材料\软著申请材料_完整版_已修改_新版.docx')

# 删除源代码中的CFD模块相关内容
# 需要删除的源代码文件标识
source_files_to_delete = [
    '文件：前处理',
    'CFD_module',
    'pre_processor_window',
    'file_handler',
    'pre_processor_config',
    'fan_id_generator',
    'scene_generator',
    'save_calculation_file',
    'load_parameters',
    'save_parameters',
    'fan_boundary_conditions',
]

# 删除包含这些内容的段落
paragraphs_to_delete = []
for i, para in enumerate(doc.paragraphs):
    text = para.text
    for marker in source_files_to_delete:
        if marker in text:
            paragraphs_to_delete.append(i)
            break

# 删除段落
for idx in sorted(set(paragraphs_to_delete), reverse=True):
    if idx < len(doc.paragraphs):
        p = doc.paragraphs[idx]._element
        p.getparent().remove(p)

# 保存
doc.save(r'F:\A-User\cliff\allin\软著申请材料\软著申请材料_完整版_已修改_新版.docx')
print(f'删除了源代码中的CFD模块内容，共 {len(set(paragraphs_to_delete))} 个段落')
